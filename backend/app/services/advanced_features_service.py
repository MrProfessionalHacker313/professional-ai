"""
Professional AI - Advanced Features Service
Handles: Images, Voice, Documents, Translation, Memory, Agents, Search, etc.
"""

import os
import json
import time
import uuid
import base64
import hashlib
import asyncio
import csv
import io
import re
from typing import Optional, Dict, Any, List
from pathlib import Path
from urllib.parse import urlparse
from loguru import logger
import httpx
from datetime import datetime
from cryptography.fernet import Fernet
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, desc

from app.config import settings
from app.models.advanced_features import (
    AIMemory, AIAgent, AgentExecution, Image, VoiceRecording,
    Document, Translation, WebSearch, Chatbot, ChatbotConversation,
    ScreenshotCode, CodeExplanation, MemoryType, AgentType, AgentStatus,
    ImageType, RecordingType, DocumentType, ProcessingStatus,
    LanguagePreference, HackingSession, AIProject, ScreenshotApp,
    ThreatAnalysis, VoiceCommandSession, MemoryVaultBackup, TaskBatch,
    AICourse, BusinessPlan, GeneratedFile, CompatibilityCheck,
    DeviceProfile, VoiceClone, NewsSubscription, NewsDigest, ModelRouterLog
)
from app.models.user import User
from app.models.usage import UsageLog


class AdvancedFeaturesService:
    """Service for all advanced AI features."""

    def __init__(self):
        self.encryption_key = self._get_or_create_encryption_key()
        self.cipher = Fernet(self.encryption_key)
        self.ollama_url = settings.OLLAMA_BASE_URL
        self.comfyui_url = settings.COMFYUI_URL
        self.searxng_url = settings.SEARXNG_URL
        self._http_client = httpx.AsyncClient(
            timeout=httpx.Timeout(300.0, connect=10.0),
            limits=httpx.Limits(max_connections=50, max_keepalive_connections=10),
        )

        self.available_models = {
            "text": ["llama3.1", "qwen2.5", "deepseek-r1", "mistral", "gemini-pro", "gpt-4o"],
            "code": ["llama3.1", "deepseek-r1", "qwen2.5", "gpt-4o"],
            "image": ["stable-diffusion-xl", "flux", "dall-e-3"],
            "voice": ["faster-whisper", "piper", "edge-tts"],
            "document": ["llama3.1", "qwen2.5", "gemini-pro"],
        }
        self.language_brain = {
            "ur": "Urdu", "en": "English", "hi": "Hindi", "bn": "Bengali", "pa": "Punjabi",
            "ps": "Pashto", "sd": "Sindhi", "ar": "Arabic", "fa": "Persian", "tr": "Turkish",
            "zh": "Chinese", "ja": "Japanese", "ko": "Korean", "ru": "Russian", "fr": "French",
            "de": "German", "es": "Spanish", "it": "Italian", "pt": "Portuguese", "nl": "Dutch",
            "pl": "Polish", "uk": "Ukrainian", "el": "Greek", "he": "Hebrew", "id": "Indonesian",
            "ms": "Malay", "th": "Thai", "vi": "Vietnamese", "sw": "Swahili", "tl": "Filipino",
            "ne": "Nepali", "si": "Sinhala", "ta": "Tamil", "te": "Telugu", "gu": "Gujarati",
            "mr": "Marathi", "ku": "Kurdish", "uz": "Uzbek", "kk": "Kazakh", "my": "Burmese",
            "km": "Khmer",
        }

    async def close(self):
        await self._http_client.aclose()

    async def __aenter__(self):
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        await self.close()

    def _get_or_create_encryption_key(self) -> bytes:
        """Get or create encryption key for vault."""
        from hashlib import sha256
        key = settings.ENCRYPTION_KEY
        if not key or key == "change-this-32-byte-hex-key-for-aes-256-gcm":
            key = Fernet.generate_key().decode()
            logger.warning("Using generated encryption key. Set ENCRYPTION_KEY in production!")
        else:
            key = key.encode() if isinstance(key, str) else key
            if len(key) != 44 or len(key.strip(b"-_")) != 32:
                key = sha256(key).digest()
                key = base64.urlsafe_b64encode(key).decode()
                logger.warning("Derived encryption key from ENCRYPTION_KEY. Set a valid Fernet key in production!")
        return key if isinstance(key, bytes) else key.encode()

    def _encrypt(self, data: str) -> str:
        """Encrypt data for storage."""
        return self.cipher.encrypt(data.encode()).decode()

    def _decrypt(self, encrypted_data: str) -> str:
        """Decrypt data from storage."""
        return self.cipher.decrypt(encrypted_data.encode()).decode()

    # ===================================================================
    # AI MEMORY SYSTEM
    # ===================================================================

    async def save_memory(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        memory_type: str,
        key: str,
        value: Any,
        importance: int = 5,
        metadata: Optional[Dict] = None
    ) -> AIMemory:
        """Save a memory to the user's long-term memory vault."""
        result = await db.execute(
            select(AIMemory).where(
                and_(
                    AIMemory.user_id == user_id,
                    AIMemory.memory_type == MemoryType(memory_type),
                    AIMemory.key == key
                )
            )
        )
        memory = result.scalar_one_or_none()

        value_str = json.dumps(value) if not isinstance(value, str) else value
        encrypted_value = self._encrypt(value_str)

        if memory:
            memory.value_encrypted = encrypted_value
            memory.importance_score = importance
            memory.extra_metadata = metadata or {}
            memory.access_count += 1
            memory.last_accessed_at = datetime.utcnow()
        else:
            memory = AIMemory(
                user_id=user_id,
                memory_type=MemoryType(memory_type),
                key=key,
                value_encrypted=encrypted_value,
                importance_score=importance,
                extra_metadata=metadata or {},
            )
            db.add(memory)

        await db.commit()
        await db.refresh(memory)
        return memory

    async def get_memory(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        memory_type: str,
        key: str
    ) -> Optional[Dict]:
        """Retrieve a memory from the vault."""
        result = await db.execute(
            select(AIMemory).where(
                and_(
                    AIMemory.user_id == user_id,
                    AIMemory.memory_type == MemoryType(memory_type),
                    AIMemory.key == key
                )
            )
        )
        memory = result.scalar_one_or_none()

        if memory:
            memory.access_count += 1
            memory.last_accessed_at = datetime.utcnow()
            await db.commit()

            try:
                decrypted = self._decrypt(memory.value_encrypted)
                try:
                    return json.loads(decrypted)
                except Exception:
                    return {"value": decrypted}
            except Exception:
                return {"value": memory.value_encrypted}
        return None

    async def get_user_memories(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        memory_type: Optional[str] = None,
        min_importance: int = 0
    ) -> List[Dict]:
        """Get all memories for a user."""
        query = select(AIMemory).where(AIMemory.user_id == user_id)

        if memory_type:
            query = query.where(AIMemory.memory_type == MemoryType(memory_type))

        query = query.where(AIMemory.importance_score >= min_importance)
        query = query.order_by(desc(AIMemory.importance_score), desc(AIMemory.last_accessed_at))

        result = await db.execute(query)
        memories = result.scalars().all()

        return [
            {
                "id": str(m.id),
                "type": m.memory_type.value,
                "key": m.key,
                "value": self._decrypt(m.value_encrypted),
                "importance": m.importance_score,
                "access_count": m.access_count,
                "last_accessed": m.last_accessed_at.isoformat(),
                "created_at": m.created_at.isoformat(),
            }
            for m in memories
        ]

    async def build_memory_context(self, db: AsyncSession, user_id: uuid.UUID) -> str:
        """Build context string from user's memories for AI prompts."""
        memories = await self.get_user_memories(db, user_id, min_importance=7)

        if not memories:
            return ""

        context_parts = ["USER CONTEXT & MEMORIES:"]
        for mem in memories[:10]:
            context_parts.append(f"- {mem['key']}: {mem['value']}")

        return "\n".join(context_parts)

    # ===================================================================
    # AI AGENTS
    # ===================================================================

    async def create_agent(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        name: str,
        description: str,
        agent_type: str,
        system_prompt: str,
        tools: Optional[List[str]] = None,
        config: Optional[Dict] = None
    ) -> AIAgent:
        """Create a new AI agent."""
        agent = AIAgent(
            user_id=user_id,
            name=name,
            description=description,
            agent_type=AgentType(agent_type),
            system_prompt=system_prompt,
            tools=tools or [],
            config=config or {},
        )
        db.add(agent)
        await db.commit()
        await db.refresh(agent)
        return agent

    async def execute_agent(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        agent_id: uuid.UUID,
        task_description: str,
        context: Optional[Dict] = None
    ) -> AgentExecution:
        """Execute an AI agent with multi-step reasoning."""
        result = await db.execute(
            select(AIAgent).where(
                and_(
                    AIAgent.id == agent_id,
                    AIAgent.user_id == user_id,
                    AIAgent.is_active == True
                )
            )
        )
        agent = result.scalar_one_or_none()

        if not agent:
            raise ValueError("Agent not found or inactive")

        execution = AgentExecution(
            agent_id=agent_id,
            user_id=user_id,
            task_description=task_description,
            steps=[],
            status=AgentStatus.RUNNING,
        )
        db.add(execution)
        await db.commit()
        await db.refresh(execution)

        try:
            start_time = time.time()
            steps = []

            steps.append({
                "step": 1,
                "name": "task_analysis",
                "status": "completed",
                "result": f"Analyzing task: {task_description[:100]}"
            })

            plan_prompt = f"""Agent: {agent.name}
Task: {task_description}
Context: {json.dumps(context or {})}

Create a step-by-step plan to complete this task. Output as JSON array of steps."""

            plan_result = await self._call_ollama(plan_prompt, agent.system_prompt)
            steps.append({
                "step": 2,
                "name": "planning",
                "status": "completed",
                "result": plan_result.get("content", "")[:200]
            })

            execution_prompt = f"""Execute this task following the plan:
{task_description}

Context: {json.dumps(context or {})}

Provide the complete result."""

            final_result = await self._call_ollama(execution_prompt, agent.system_prompt)
            steps.append({
                "step": 3,
                "name": "execution",
                "status": "completed",
                "result": "Task executed successfully"
            })

            execution.steps = steps
            execution.result = final_result.get("content", "")
            execution.status = AgentStatus.COMPLETED
            execution.tokens_used = final_result.get("tokens", 0)
            execution.execution_time_ms = int((time.time() - start_time) * 1000)
            execution.completed_at = datetime.utcnow()

            agent.execution_count += 1
            agent.success_rate = ((agent.success_rate * (agent.execution_count - 1) + 1) / agent.execution_count)

            await db.commit()
            await db.refresh(execution)

            await self._log_usage(db, user_id, "agent_execution", execution.tokens_used, execution.execution_time_ms)

            return execution

        except Exception as e:
            execution.status = AgentStatus.FAILED
            execution.error_message = str(e)
            execution.completed_at = datetime.utcnow()

            agent.execution_count += 1
            agent.success_rate = ((agent.success_rate * (agent.execution_count - 1)) / agent.execution_count)

            await db.commit()
            raise

    # ===================================================================
    # IMAGE GENERATION & ANALYSIS
    # ===================================================================

    async def generate_image(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        prompt: str,
        negative_prompt: str = "",
        model: str = "stable-diffusion-xl",
        width: int = 1024,
        height: int = 1024,
        steps: int = 30
    ) -> Dict:
        """Generate an image using Stable Diffusion or Flux."""
        start_time = time.time()

        workflow = self._build_comfyui_workflow(prompt, negative_prompt, model, width, height, steps)

        async with httpx.AsyncClient(timeout=300.0) as client:
            response = await client.post(f"{self.comfyui_url}/prompt", json=workflow)
            response.raise_for_status()
            prompt_id = response.json().get("prompt_id")

            result = await self._poll_comfyui_result(client, prompt_id)

        execution_time = int((time.time() - start_time) * 1000)

        image = Image(
            user_id=user_id,
            image_type=ImageType.GENERATED,
            storage_path=result.get("image_path", ""),
            thumbnail_path=result.get("thumbnail_path"),
            prompt=prompt,
            negative_prompt=negative_prompt,
            model_used=model,
            parameters={"width": width, "height": height, "steps": steps},
            width=width,
            height=height,
            file_size_bytes=result.get("file_size", 0),
            mime_type="image/png",
        )
        db.add(image)
        await db.commit()
        await db.refresh(image)

        await self._log_usage(db, user_id, "image_generation", 0, execution_time)

        return {
            "id": str(image.id),
            "image_path": image.storage_path,
            "thumbnail_path": image.thumbnail_path,
            "prompt": prompt,
            "model": model,
            "execution_time_ms": execution_time,
        }

    def _build_comfyui_workflow(self, prompt: str, negative_prompt: str, model: str, width: int, height: int, steps: int) -> Dict:
        """Build ComfyUI workflow JSON."""
        return {
            "3": {
                "class_type": "KSampler",
                "inputs": {
                    "seed": int(time.time()),
                    "steps": steps,
                    "cfg": 7.5,
                    "sampler_name": "DPM++ 2M Karras",
                    "scheduler": "karras",
                    "denoise": 1,
                    "model": ["4", 0],
                    "positive": ["6", 0],
                    "negative": ["7", 0],
                    "latent_image": ["5", 0],
                }
            },
            "4": {
                "class_type": "CheckpointLoaderSimple",
                "inputs": {
                    "ckpt_name": model + ".safetensors"
                }
            },
            "5": {
                "class_type": "EmptyLatentImage",
                "inputs": {
                    "width": width,
                    "height": height,
                    "batch_size": 1
                }
            },
            "6": {
                "class_type": "CLIPTextEncode",
                "inputs": {
                    "text": prompt,
                    "clip": ["4", 1]
                }
            },
            "7": {
                "class_type": "CLIPTextEncode",
                "inputs": {
                    "text": negative_prompt,
                    "clip": ["4", 1]
                }
            },
            "8": {
                "class_type": "VAEDecode",
                "inputs": {
                    "samples": ["3", 0],
                    "vae": ["4", 2]
                }
            },
            "9": {
                "class_type": "SaveImage",
                "inputs": {
                    "filename_prefix": "ProAI",
                    "images": ["8", 0]
                }
            }
        }

    async def _poll_comfyui_result(self, client: httpx.AsyncClient, prompt_id: str, max_wait: int = 300) -> Dict:
        """Poll ComfyUI for workflow completion."""
        for _ in range(max_wait):
            response = await client.get(f"{self.comfyui_url}/history/{prompt_id}")
            if response.status_code == 200:
                history = response.json()
                if prompt_id in history:
                    status = history[prompt_id].get("status", {})
                    if status.get("completed", False):
                        outputs = history[prompt_id].get("outputs", {})
                        for node_id, node_output in outputs.items():
                            if "images" in node_output:
                                images = node_output["images"]
                                if images:
                                    img = images[0]
                                    return {
                                        "image_path": f"{self.comfyui_url}/view?filename={img['filename']}",
                                        "file_size": img.get("size", 0),
                                    }
            await asyncio.sleep(1)

        raise TimeoutError("Image generation timed out")

    async def analyze_image(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        image_path: str,
        analysis_type: str = "describe"
    ) -> Dict:
        """Analyze an image - describe, OCR, or edit."""
        start_time = time.time()

        with open(image_path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode()

        if analysis_type == "describe":
            prompt = "Describe this image in detail. What do you see?"
            result = await self._call_ollama_vision(prompt, image_data)
            analysis_text = result.get("content", "")
        elif analysis_type == "ocr":
            analysis_text = await self._extract_text_from_image(image_path)
        else:
            analysis_text = "Analysis type not supported"

        execution_time = int((time.time() - start_time) * 1000)

        image = Image(
            user_id=user_id,
            image_type=ImageType.ANALYZED,
            storage_path=image_path,
            prompt=analysis_type,
            model_used="vision-model",
        )
        db.add(image)
        await db.commit()

        await self._log_usage(db, user_id, "image_analysis", 0, execution_time)

        return {
            "analysis": analysis_text,
            "type": analysis_type,
            "execution_time_ms": execution_time,
        }

    async def _extract_text_from_image(self, image_path: str) -> str:
        """Extract text from image using OCR."""
        return "OCR text extraction would happen here using Tesseract"

    async def _call_ollama_vision(self, prompt: str, image_data: str) -> Dict:
        """Call Ollama with vision model."""
        url = f"{self.ollama_url}/api/generate"

        payload = {
            "model": "llava",
            "prompt": prompt,
            "images": [image_data],
            "stream": False,
        }

        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(url, json=payload)
            response.raise_for_status()
            result = response.json()

        return {
            "content": result.get("response", ""),
            "tokens": result.get("eval_count", 0),
        }

    # ===================================================================
    # VOICE FEATURES
    # ===================================================================

    async def speech_to_text(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        audio_path: str,
        language: str = "en"
    ) -> Dict:
        """Convert speech to text using faster-whisper."""
        start_time = time.time()

        async with httpx.AsyncClient(timeout=120.0) as client:
            with open(audio_path, "rb") as f:
                files = {"file": f}
                data = {"language": language}
                response = await client.post(
                    f"{settings.WHISPER_API_URL}/transcribe",
                    files=files,
                    data=data
                )
                response.raise_for_status()
                result = response.json()

        execution_time = int((time.time() - start_time) * 1000)

        recording = VoiceRecording(
            user_id=user_id,
            recording_type=RecordingType.INPUT,
            storage_path=audio_path,
            language=language,
            transcription=result.get("text", ""),
            model_used="faster-whisper",
        )
        db.add(recording)
        await db.commit()

        await self._log_usage(db, user_id, "voice_input", 0, execution_time)

        return {
            "text": result.get("text", ""),
            "language": language,
            "execution_time_ms": execution_time,
        }

    async def text_to_speech(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        text: str,
        language: str = "en",
        voice: str = "default"
    ) -> Dict:
        """Convert text to speech using Piper TTS."""
        start_time = time.time()

        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                f"{settings.TTS_API_URL}/synthesize",
                json={
                    "text": text,
                    "language": language,
                    "voice": voice,
                }
            )
            response.raise_for_status()
            result = response.json()

        execution_time = int((time.time() - start_time) * 1000)

        recording = VoiceRecording(
            user_id=user_id,
            recording_type=RecordingType.OUTPUT,
            storage_path=result.get("audio_path", ""),
            language=language,
            transcription=text,
            model_used="piper-tts",
        )
        db.add(recording)
        await db.commit()

        await self._log_usage(db, user_id, "voice_output", 0, execution_time)

        return {
            "audio_path": result.get("audio_path"),
            "text": text,
            "language": language,
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # DOCUMENT ANALYSIS
    # ===================================================================

    async def upload_document(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        file_path: str,
        original_filename: str
    ) -> Document:
        """Upload and process a document."""
        file_size = os.path.getsize(file_path)
        mime_type = self._get_mime_type(file_path)
        doc_type = self._get_document_type(file_path)

        document = Document(
            user_id=user_id,
            document_type=doc_type,
            original_filename=original_filename,
            storage_path=file_path,
            file_size_bytes=file_size,
            mime_type=mime_type,
            processing_status=ProcessingStatus.PENDING,
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)

        asyncio.create_task(self._process_document(db, document.id))

        return document

    async def _process_document(self, db: AsyncSession, document_id: uuid.UUID):
        """Process uploaded document (extract text, summarize, etc.)."""
        async with db.begin():
            result = await db.execute(select(Document).where(Document.id == document_id).with_for_update())
            document = result.scalar_one_or_none()

            if not document:
                return

            try:
                document.processing_status = ProcessingStatus.PROCESSING
                await db.commit()

                if document.document_type == DocumentType.PDF:
                    text = await self._extract_pdf_text(document.storage_path)
                elif document.document_type == DocumentType.DOCX:
                    text = await self._extract_docx_text(document.storage_path)
                else:
                    text = await self._extract_text(document.storage_path)

                document.extracted_text = text
                document.word_count = len(text.split())
                document.language_detected = self._detect_language(text)

                summary_prompt = f"Summarize this document concisely:\n\n{text[:4000]}"
                summary_result = await self._call_ollama(summary_prompt)
                document.summary = summary_result.get("content", "")

                document.processing_status = ProcessingStatus.COMPLETED
                document.processed_at = datetime.utcnow()
                await db.commit()

            except Exception as e:
                document.processing_status = ProcessingStatus.FAILED
                await db.commit()
                logger.error(f"Document processing failed: {str(e)}")

    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF."""
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    async def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    def _detect_language(self, text: str) -> str:
        """Detect language of text."""
        try:
            from langdetect import detect
            return detect(text)
        except Exception:
            return "en"

    def _normalize_lang_code(self, lang_code: str) -> str:
        """Normalize language code to a primary subtag."""
        if not lang_code:
            return "en"
        return lang_code.strip().lower().split("-")[0]

    def _safe_json_loads(self, raw: str, default: Dict[str, Any]) -> Dict[str, Any]:
        try:
            parsed = json.loads(raw)
            return parsed if isinstance(parsed, dict) else default
        except Exception:
            return default

    def _get_mime_type(self, file_path: str) -> str:
        """Get MIME type from file extension."""
        ext = Path(file_path).suffix.lower()
        mime_types = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".txt": "text/plain",
            ".png": "image/png",
            ".jpg": "image/jpeg",
        }
        return mime_types.get(ext, "application/octet-stream")

    def _get_document_type(self, file_path: str) -> DocumentType:
        """Get document type from file extension."""
        ext = Path(file_path).suffix.lower()
        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".png": DocumentType.IMAGE,
            ".jpg": DocumentType.IMAGE,
        }
        return type_map.get(ext, DocumentType.TXT)

    # ===================================================================
    # TRANSLATION
    # ===================================================================

    async def translate_text(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        text: str,
        source_lang: str,
        target_lang: str,
        context_type: str = "chat",
        context_id: Optional[uuid.UUID] = None
    ) -> Dict:
        """Translate text between languages."""
        start_time = time.time()

        prompt = f"Translate the following text from {source_lang} to {target_lang}.\n\nText: {text}\n\nTranslation:"
        result = await self._call_ollama(prompt)

        translated_text = result.get("content", text)
        execution_time = int((time.time() - start_time) * 1000)

        translation = Translation(
            user_id=user_id,
            source_language=source_lang,
            target_language=target_lang,
            original_text=text,
            translated_text=translated_text,
            context_type=context_type,
            context_id=context_id,
            model_used="llama3.1",
        )
        db.add(translation)
        await db.commit()

        await self._log_usage(db, user_id, "translation", 0, execution_time)

        return {
            "original_text": text,
            "translated_text": translated_text,
            "source_lang": source_lang,
            "target_lang": target_lang,
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # WEB SEARCH
    # ===================================================================

    async def web_search(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        query: str,
        search_engine: str = "searxng"
    ) -> Dict:
        """Perform web search using SearXNG or Serper."""
        start_time = time.time()

        if search_engine == "searxng":
            results = await self._search_searxng(query)
        elif search_engine == "serper":
            results = await self._search_serper(query)
        else:
            results = {"error": "Unsupported search engine"}

        execution_time = int((time.time() - start_time) * 1000)

        web_search = WebSearch(
            user_id=user_id,
            query=query,
            search_engine=search_engine,
            results=results,
            result_count=len(results.get("results", [])),
            execution_time_ms=execution_time,
        )
        db.add(web_search)
        await db.commit()

        await self._log_usage(db, user_id, "web_search", 0, execution_time)

        return {
            "query": query,
            "results": results.get("results", []),
            "result_count": len(results.get("results", [])),
            "execution_time_ms": execution_time,
        }

    async def _search_searxng(self, query: str) -> Dict:
        """Search using self-hosted SearXNG."""
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(
                f"{self.searxng_url}/search",
                params={"q": query, "format": "json"},
            )
            response.raise_for_status()
            return response.json()

    async def _search_serper(self, query: str) -> Dict:
        """Search using Serper API."""
        api_key = os.getenv("SERPER_API_KEY")
        if not api_key:
            return {"error": "Serper API key not configured"}

        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                "https://google.serper.dev/search",
                headers={"X-API-KEY": api_key},
                json={"q": query},
            )
            response.raise_for_status()
            return response.json()

    # ===================================================================
    # CODE EXPLAINER
    # ===================================================================

    async def explain_code(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        code: str,
        language: str,
        user_language: str = "en"
    ) -> Dict:
        """Explain code line-by-line in user's language."""
        start_time = time.time()

        prompt = f"""Explain this {language} code line-by-line in {user_language}.

Code:
```{language}
{code}
```

Provide:
1. Overall purpose
2. Line-by-line explanation
3. Key concepts used
4. Potential improvements

Format as JSON with keys: overview, line_by_line (array), concepts, improvements"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            explanation_data = json.loads(result.get("content", "{}"))
        except Exception:
            explanation_data = {"overview": result.get("content", "")}

        code_exp = CodeExplanation(
            user_id=user_id,
            original_code=code,
            language=language,
            explanation=explanation_data.get("overview", ""),
            line_by_line=explanation_data.get("line_by_line", []),
            model_used="llama3.1",
        )
        db.add(code_exp)
        await db.commit()

        await self._log_usage(db, user_id, "code_explanation", 0, execution_time)

        return {
            "overview": explanation_data.get("overview", ""),
            "line_by_line": explanation_data.get("line_by_line", []),
            "concepts": explanation_data.get("concepts", []),
            "improvements": explanation_data.get("improvements", []),
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # SCREENSHOT TO CODE
    # ===================================================================

    async def screenshot_to_code(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        image_path: str,
        framework: str = "html"
    ) -> Dict:
        """Convert screenshot to HTML/CSS code."""
        start_time = time.time()

        with open(image_path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode()

        prompt = f"""Convert this screenshot to {framework} code.
Generate clean, semantic, responsive HTML/CSS code that replicates this design.
Include all visual elements, colors, spacing, and layout.
Output only the code, no explanations."""

        result = await self._call_ollama_vision(prompt, image_data)
        execution_time = int((time.time() - start_time) * 1000)

        image = Image(
            user_id=user_id,
            image_type=ImageType.ANALYZED,
            storage_path=image_path,
            prompt="screenshot_to_code",
        )
        db.add(image)
        await db.commit()
        await db.refresh(image)

        screenshot_code = ScreenshotCode(
            user_id=user_id,
            image_id=image.id,
            generated_code=result.get("content", ""),
            framework=framework,
            language="html",
            model_used="vision-model",
        )
        db.add(screenshot_code)
        await db.commit()

        await self._log_usage(db, user_id, "screenshot_to_code", 0, execution_time)

        return {
            "code": result.get("content", ""),
            "framework": framework,
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # CHATBOT BUILDER
    # ===================================================================

    async def create_chatbot(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        name: str,
        description: str,
        system_prompt: str,
        welcome_message: Optional[str] = None,
        suggested_prompts: Optional[List[str]] = None
    ) -> Chatbot:
        """Create a custom chatbot."""
        chatbot = Chatbot(
            user_id=user_id,
            name=name,
            description=description,
            system_prompt=system_prompt,
            welcome_message=welcome_message or f"Hello! I'm {name}. How can I help you?",
            suggested_prompts=suggested_prompts or [],
        )
        db.add(chatbot)
        await db.commit()
        await db.refresh(chatbot)
        return chatbot

    async def chat_with_bot(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        chatbot_id: uuid.UUID,
        message: str,
        session_id: Optional[str] = None
    ) -> Dict:
        """Chat with a custom chatbot."""
        result = await db.execute(select(Chatbot).where(
            Chatbot.id == chatbot_id,
            Chatbot.user_id == user_id
        ))
        chatbot = result.scalar_one_or_none()

        if not chatbot:
            raise ValueError("Chatbot not found or access denied")

        if not session_id:
            session_id = str(uuid.uuid4())

        conv_result = await db.execute(
            select(ChatbotConversation).where(
                and_(
                    ChatbotConversation.chatbot_id == chatbot_id,
                    ChatbotConversation.session_id == session_id
                )
            )
        )
        conversation = conv_result.scalar_one_or_none()

        if not conversation:
            conversation = ChatbotConversation(
                chatbot_id=chatbot_id,
                user_id=user_id,
                session_id=session_id,
                messages=[],
            )
            db.add(conversation)

        messages = conversation.messages or []
        messages.append({"role": "user", "content": message})

        ai_result = await self._call_ollama(message, chatbot.system_prompt)
        ai_response = ai_result.get("content", "")

        messages.append({"role": "assistant", "content": ai_response})
        conversation.messages = messages
        conversation.last_message_at = datetime.utcnow()

        chatbot.conversation_count += 1

        await db.commit()

        return {
            "response": ai_response,
            "session_id": session_id,
            "message_count": len(messages),
        }

    # ===================================================================
    # MODEL ROUTER
    # ===================================================================

    async def route_task(
        self,
        task_type: str,
        task_description: str,
        user_preferences: Optional[Dict] = None
    ) -> Dict:
        """Route task to the best model based on task type and performance."""
        model_map = {
            "text": "llama3.1",
            "code": "deepseek-r1",
            "image": "stable-diffusion-xl",
            "voice": "faster-whisper",
            "document": "qwen2.5",
            "search": "llama3.1",
        }

        selected_model = model_map.get(task_type, "llama3.1")
        provider = "ollama" if selected_model in ["llama3.1", "qwen2.5", "deepseek-r1"] else "cloud"

        return {
            "model": selected_model,
            "provider": provider,
            "task_type": task_type,
            "reason": f"Best model for {task_type} tasks based on performance metrics",
        }

    # ===================================================================
    # FEATURE 1: 40+ LANGUAGE BRAIN
    # ===================================================================

    async def detect_user_language(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        text: str
    ) -> Dict:
        """Detect user language from text."""
        start_time = time.time()

        detected = self._normalize_lang_code(self._detect_language(text))

        lang_pref_result = await db.execute(
            select(LanguagePreference).where(LanguagePreference.user_id == user_id)
        )
        lang_pref = lang_pref_result.scalar_one_or_none()

        if not lang_pref:
            lang_pref = LanguagePreference(
                user_id=user_id,
                detected_language=detected,
                preferred_language=detected if detected in self.language_brain else "en",
            )
            db.add(lang_pref)
        else:
            lang_pref.detected_language = detected
            if detected in self.language_brain:
                lang_pref.preferred_language = detected
            lang_pref.confidence_score = 0.9 if len(text) > 10 else 0.7

        await db.commit()

        execution_time = int((time.time() - start_time) * 1000)
        await self._log_usage(db, user_id, "language_detection", 0, execution_time)

        return {
            "detected_language": detected,
            "detected_language_name": self.language_brain.get(detected, "Unknown"),
            "supported": detected in self.language_brain,
            "confidence_score": lang_pref.confidence_score,
            "execution_time_ms": execution_time,
        }

    async def auto_translate_to_user_language(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        text: str,
        target_lang: str
    ) -> Dict:
        """Auto-translate text to user's preferred language."""
        start_time = time.time()

        detected = self._normalize_lang_code(self._detect_language(text))
        normalized_target = self._normalize_lang_code(target_lang)

        if detected == normalized_target:
            return {
                "original_text": text,
                "translated_text": text,
                "source_lang": detected,
                "target_lang": normalized_target,
                "execution_time_ms": 0,
            }

        prompt = (
            f"Translate the following text from {detected} to {normalized_target}.\n\n"
            f"Text: {text}\n\n"
            "Return natural phrasing with local idioms while preserving meaning.\n\nTranslation:"
        )
        result = await self._call_ollama(prompt)
        translated = result.get("content", text)
        execution_time = int((time.time() - start_time) * 1000)

        translation = Translation(
            user_id=user_id,
            source_language=detected,
            target_language=normalized_target,
            original_text=text,
            translated_text=translated,
            context_type="auto_language_brain",
            model_used="llama3.1",
        )
        db.add(translation)
        await db.commit()

        await self._log_usage(db, user_id, "auto_translate", 0, execution_time)

        return {
            "original_text": text,
            "translated_text": translated,
            "source_lang": detected,
            "target_lang": normalized_target,
            "execution_time_ms": execution_time,
        }

    async def generate_native_language_reply(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        user_text: str,
        task: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Generate a reply in the user's detected/preferred language."""
        profile_result = await db.execute(
            select(LanguagePreference).where(LanguagePreference.user_id == user_id)
        )
        profile = profile_result.scalar_one_or_none()

        detected = self._normalize_lang_code(self._detect_language(user_text))
        target_lang = detected
        if profile and profile.preferred_language:
            target_lang = self._normalize_lang_code(profile.preferred_language)
        if target_lang not in self.language_brain:
            target_lang = detected if detected in self.language_brain else "en"

        prompt = (
            f"User message: {user_text}\n"
            f"Task context: {task or 'general assistant'}\n\n"
            f"Reply natively in {self.language_brain.get(target_lang, 'English')} ({target_lang}). "
            "Keep the answer natural, culturally fluent, and concise. "
            "Use local idioms/slang only when appropriate and safe."
        )
        result = await self._call_ollama(prompt)
        return {
            "language": target_lang,
            "language_name": self.language_brain.get(target_lang, "English"),
            "reply": result.get("content", ""),
        }

    async def get_language_brain_context(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> str:
        """Get language brain context for AI prompts."""
        result = await db.execute(
            select(LanguagePreference).where(LanguagePreference.user_id == user_id)
        )
        lang_pref = result.scalar_one_or_none()

        if not lang_pref:
            return "User language preference: English (default)"

        parts = [f"User language: {lang_pref.detected_language or lang_pref.preferred_language}"]
        if lang_pref.language_context:
            for k, v in list(lang_pref.language_context.items())[:5]:
                parts.append(f"- {k}: {v}")

        return "\n".join(parts)

    # ===================================================================
    # FEATURE 2: LIVE HACKING LAB
    # ===================================================================

    async def create_hacking_session(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        attack_type: str,
        target_description: str
    ) -> Dict:
        """Create a new safe hacking session."""
        allowed_types = {"sqli", "xss", "brute_force"}
        normalized_attack = attack_type.strip().lower().replace(" ", "_")
        if normalized_attack not in allowed_types:
            raise ValueError("Only educational lab types are allowed: sqli, xss, brute_force")

        session = HackingSession(
            user_id=user_id,
            attack_type=normalized_attack,
            target_description=target_description,
            total_steps=5,
            steps=[],
            risk_level="safe",
        )
        db.add(session)
        await db.commit()
        await db.refresh(session)

        await self._log_usage(db, user_id, "hacking_session_create", 0, 0)

        return {
            "id": str(session.id),
            "attack_type": session.attack_type,
            "target_description": session.target_description,
            "status": session.status.value,
            "current_step": session.current_step,
            "total_steps": session.total_steps,
            "risk_level": session.risk_level,
            "created_at": session.created_at.isoformat(),
        }

    async def run_safe_attack(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        session_id: Any,
        attack_step: str,
        payload: str
    ) -> Dict:
        """Run a safe attack step with AI feedback."""
        try:
            session_uuid = session_id if isinstance(session_id, uuid.UUID) else uuid.UUID(str(session_id))
        except Exception:
            raise ValueError("Invalid session_id")

        result = await db.execute(
            select(HackingSession).where(
                and_(
                    HackingSession.id == session_uuid,
                    HackingSession.user_id == user_id
                )
            )
        )
        session = result.scalar_one_or_none()

        if not session:
            raise ValueError("Hacking session not found")

        start_time = time.time()
        prompt = f"""You are running a SAFE, educational local-only cybersecurity lab simulation.
    Never provide real-world attack instructions or deployable exploit chains.
    Focus on detection, mitigation, defensive engineering, and sandbox-only learning.
Attack Type: {session.attack_type}
Target: {session.target_description}
Step: {attack_step}
Payload: {payload}

Analyze this step and provide:
1. Whether this would succeed in a real scenario
2. What defenses would block it
3. Educational feedback
4. Next recommended step

Output as JSON: {{"success": bool, "defenses": [str], "feedback": str, "next_step": str}}"""
        ai_result = await self._call_ollama(prompt)

        execution_time = int((time.time() - start_time) * 1000)

        step_data = {
            "step": session.current_step + 1,
            "name": attack_step,
            "payload": payload,
            "result": ai_result.get("content", "")[:500],
            "execution_time_ms": execution_time,
        }

        steps = session.steps or []
        steps.append(step_data)
        session.steps = steps
        session.current_step = len(steps)
        session.ai_feedback = ai_result.get("content", "")

        if session.current_step >= session.total_steps:
            session.status = AgentStatus.COMPLETED
            session.completed_at = datetime.utcnow()

        await db.commit()

        await self._log_usage(db, user_id, "hacking_attack_step", 0, execution_time)

        feedback = self._safe_json_loads(ai_result.get("content", "{}"), {"feedback": ai_result.get("content", "")})

        return {
            "session_id": str(session.id),
            "step": step_data["step"],
            "success": feedback.get("success", True),
            "defenses": feedback.get("defenses", []),
            "feedback": feedback.get("feedback", ""),
            "next_step": feedback.get("next_step", ""),
            "execution_time_ms": execution_time,
        }

    async def get_hacking_session_progress(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        session_id: Any
    ) -> Dict:
        """Get progress of a hacking session."""
        try:
            session_uuid = session_id if isinstance(session_id, uuid.UUID) else uuid.UUID(str(session_id))
        except Exception:
            raise ValueError("Invalid session_id")

        result = await db.execute(
            select(HackingSession).where(
                and_(
                    HackingSession.id == session_uuid,
                    HackingSession.user_id == user_id
                )
            )
        )
        session = result.scalar_one_or_none()

        if not session:
            raise ValueError("Hacking session not found")

        return {
            "id": str(session.id),
            "attack_type": session.attack_type,
            "status": session.status.value,
            "current_step": session.current_step,
            "total_steps": session.total_steps,
            "progress_percent": round((session.current_step / session.total_steps) * 100, 1),
            "steps": session.steps or [],
            "risk_level": session.risk_level,
            "created_at": session.created_at.isoformat(),
        }

    async def list_hacking_sessions(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> List[Dict]:
        """List all hacking sessions for a user."""
        result = await db.execute(
            select(HackingSession).where(HackingSession.user_id == user_id)
            .order_by(desc(HackingSession.created_at))
        )
        sessions = result.scalars().all()

        return [
            {
                "id": str(s.id),
                "attack_type": s.attack_type,
                "status": s.status.value,
                "current_step": s.current_step,
                "total_steps": s.total_steps,
                "risk_level": s.risk_level,
                "created_at": s.created_at.isoformat(),
            }
            for s in sessions
        ]

    # ===================================================================
    # FEATURE 3: AI PROJECT ASSISTANT
    # ===================================================================

    async def build_project(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        description: str,
        stack: List[str]
    ) -> Dict:
        """Build a project from description and stack."""
        start_time = time.time()

        prompt = f"""Generate a project structure for the following:
Description: {description}
Tech Stack: {', '.join(stack)}

Output JSON with this structure:
{{
    "name": "project_name",
    "files": {{
        "main.py": "code here",
        "requirements.txt": "packages here",
        "README.md": "readme here"
    }},
    "structure": ["main.py", "requirements.txt", "README.md"],
    "setup_instructions": ["step1", "step2"]
}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            project_data = json.loads(result.get("content", "{}"))
        except Exception:
            project_data = {
                "name": "ai_generated_project",
                "files": {"main.py": result.get("content", "")},
                "structure": ["main.py"],
                "setup_instructions": [],
            }

        project = AIProject(
            user_id=user_id,
            name=project_data.get("name", "AI Project"),
            description=description,
            stack=stack,
            files=project_data.get("files", {}),
            status="completed",
            progress_percent=100,
            model_used="llama3.1",
        )
        db.add(project)
        await db.commit()
        await db.refresh(project)

        await self._log_usage(db, user_id, "project_build", 0, execution_time)

        return {
            "id": str(project.id),
            "name": project.name,
            "description": project.description,
            "stack": project.stack,
            "files": project.files,
            "structure": project_data.get("structure", []),
            "setup_instructions": project_data.get("setup_instructions", []),
            "status": project.status,
            "execution_time_ms": execution_time,
        }

    async def get_project(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        project_id: uuid.UUID
    ) -> Dict:
        """Get project details."""
        result = await db.execute(
            select(AIProject).where(
                and_(
                    AIProject.id == project_id,
                    AIProject.user_id == user_id
                )
            )
        )
        project = result.scalar_one_or_none()

        if not project:
            raise ValueError("Project not found")

        return {
            "id": str(project.id),
            "name": project.name,
            "description": project.description,
            "stack": project.stack,
            "files": project.files,
            "status": project.status,
            "progress_percent": project.progress_percent,
            "build_log": project.build_log,
            "created_at": project.created_at.isoformat(),
        }

    # ===================================================================
    # FEATURE 4: SCREENSHOT/PHOTO → FULL APP
    # ===================================================================

    async def generate_full_app_from_screenshot(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        image_path: str,
        platform: str,
        framework: str,
        include_api: bool = True,
        include_auth: bool = False,
        styling: str = "tailwind",
    ) -> Dict:
        """Generate a full app from a screenshot."""
        start_time = time.time()

        with open(image_path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode()

        prompt = f"""Generate a complete {platform} app using {framework} from this screenshot.
    Styling system: {styling}
    Include backend API: {include_api}
    Include authentication: {include_auth}
Output JSON:
{{
    "app_name": "name",
    "files": {{
        "main.py": "code",
        "config.yaml": "config"
    }},
    "structure": ["main.py", "config.yaml"],
    "dependencies": ["package1", "package2"],
    "run_instructions": ["step1", "step2"]
}}"""

        result = await self._call_ollama_vision(prompt, image_data)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            app_data = json.loads(result.get("content", "{}"))
        except Exception:
            app_data = {
                "app_name": "screenshot_app",
                "files": {"main.py": result.get("content", "")},
                "structure": ["main.py"],
                "dependencies": [],
                "run_instructions": [],
            }

        image = Image(
            user_id=user_id,
            image_type=ImageType.ANALYZED,
            storage_path=image_path,
            prompt="screenshot_to_full_app",
        )
        db.add(image)
        await db.commit()
        await db.refresh(image)

        screenshot_app = ScreenshotApp(
            user_id=user_id,
            image_id=image.id,
            platform=platform,
            framework=framework,
            app_files=app_data.get("files", {}),
            app_structure={
                "structure": app_data.get("structure", []),
                "dependencies": app_data.get("dependencies", []),
                "run_instructions": app_data.get("run_instructions", []),
            },
            model_used="vision-model",
        )
        db.add(screenshot_app)
        await db.commit()

        await self._log_usage(db, user_id, "screenshot_to_full_app", 0, execution_time)

        return {
            "id": str(screenshot_app.id),
            "app_name": app_data.get("app_name", "generated_app"),
            "platform": platform,
            "framework": framework,
            "files": screenshot_app.app_files,
            "structure": screenshot_app.app_structure,
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # FEATURE 5: AI DETECTIVE
    # ===================================================================

    async def analyze_file(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        file_path: str,
        file_type: str
    ) -> Dict:
        """Analyze file for threats."""
        start_time = time.time()

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()[:5000]
        except Exception:
            content = ""

        prompt = f"""Analyze this {file_type} file for security threats, malware indicators, and suspicious patterns.
File content (first 5000 chars):
{content}

Output JSON: {{"threat_level": "safe|low|medium|high|critical", "threats": [{{"type": str, "severity": str, "description": str}}], "recommendations": [str], "confidence_score": float}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            analysis = json.loads(result.get("content", "{}"))
        except Exception:
            analysis = {
                "threat_level": "unknown",
                "threats": [],
                "recommendations": ["Unable to parse analysis"],
                "confidence_score": 0.0,
            }

        threat = ThreatAnalysis(
            user_id=user_id,
            analysis_type="file",
            target=file_path,
            threat_level=analysis.get("threat_level", "unknown"),
            threats_found=analysis.get("threats", []),
            recommendations=analysis.get("recommendations", []),
            confidence_score=analysis.get("confidence_score", 0.0),
            model_used="llama3.1",
        )
        db.add(threat)
        await db.commit()

        await self._log_usage(db, user_id, "file_analysis", 0, execution_time)

        return {
            "id": str(threat.id),
            "file_path": file_path,
            "file_type": file_type,
            "threat_level": threat.threat_level,
            "threats": threat.threats_found,
            "recommendations": threat.recommendations,
            "confidence_score": threat.confidence_score,
            "execution_time_ms": execution_time,
        }

    async def analyze_link(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        url: str
    ) -> Dict:
        """Analyze link for safety."""
        start_time = time.time()

        parsed = urlparse(url)
        host = (parsed.hostname or "").lower()
        local_indicators = []
        if "xn--" in host:
            local_indicators.append("punycode_domain")
        if re.fullmatch(r"\d+\.\d+\.\d+\.\d+", host or ""):
            local_indicators.append("raw_ip_hostname")
        if any(x in url.lower() for x in ["verify-account", "reset-password", "urgent", "gift-card"]):
            local_indicators.append("social_engineering_keywords")

        prompt = f"""Analyze this URL for safety threats: {url}
Local indicators found: {', '.join(local_indicators) if local_indicators else 'none'}
Output JSON: {{"threat_level": "safe|low|medium|high|critical", "threats": [str], "recommendations": [str], "confidence_score": float}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        analysis = self._safe_json_loads(result.get("content", "{}"), {
            "threat_level": "unknown",
            "threats": [],
            "recommendations": ["Unable to parse analysis"],
            "confidence_score": 0.0,
        })
        if local_indicators:
            analysis["threats"] = list(analysis.get("threats", [])) + local_indicators
            if analysis.get("threat_level") in ("safe", "low", None, ""):
                analysis["threat_level"] = "medium"

        threat = ThreatAnalysis(
            user_id=user_id,
            analysis_type="link",
            target=url,
            threat_level=analysis.get("threat_level", "unknown"),
            threats_found=analysis.get("threats", []),
            recommendations=analysis.get("recommendations", []),
            confidence_score=analysis.get("confidence_score", 0.0),
            model_used="llama3.1",
        )
        db.add(threat)
        await db.commit()

        await self._log_usage(db, user_id, "link_analysis", 0, execution_time)

        return {
            "id": str(threat.id),
            "url": url,
            "threat_level": threat.threat_level,
            "threats": threat.threats_found,
            "recommendations": threat.recommendations,
            "confidence_score": threat.confidence_score,
            "execution_time_ms": execution_time,
        }

    async def analyze_email(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        email_content: str
    ) -> Dict:
        """Analyze email for phishing and threats."""
        start_time = time.time()

        prompt = f"""Analyze this email for phishing, scams, and security threats:
{email_content[:3000]}

Output JSON: {{"threat_level": "safe|low|medium|high|critical", "threats": [str], "recommendations": [str], "confidence_score": float}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            analysis = json.loads(result.get("content", "{}"))
        except Exception:
            analysis = {
                "threat_level": "unknown",
                "threats": [],
                "recommendations": ["Unable to parse analysis"],
                "confidence_score": 0.0,
            }

        threat = ThreatAnalysis(
            user_id=user_id,
            analysis_type="email",
            target=email_content[:200],
            threat_level=analysis.get("threat_level", "unknown"),
            threats_found=analysis.get("threats", []),
            recommendations=analysis.get("recommendations", []),
            confidence_score=analysis.get("confidence_score", 0.0),
            model_used="llama3.1",
        )
        db.add(threat)
        await db.commit()

        await self._log_usage(db, user_id, "email_analysis", 0, execution_time)

        return {
            "id": str(threat.id),
            "threat_level": threat.threat_level,
            "threats": threat.threats_found,
            "recommendations": threat.recommendations,
            "confidence_score": threat.confidence_score,
            "execution_time_ms": execution_time,
        }

    async def get_detective_history(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> List[Dict]:
        """Get detective analysis history."""
        result = await db.execute(
            select(ThreatAnalysis).where(ThreatAnalysis.user_id == user_id)
            .order_by(desc(ThreatAnalysis.created_at))
        )
        analyses = result.scalars().all()

        return [
            {
                "id": str(a.id),
                "analysis_type": a.analysis_type,
                "target": a.target[:100] if a.target else "",
                "threat_level": a.threat_level,
                "threats_count": len(a.threats_found),
                "confidence_score": a.confidence_score,
                "created_at": a.created_at.isoformat(),
            }
            for a in analyses
        ]

    # ===================================================================
    # FEATURE 6: VOICE COMMAND MODE
    # ===================================================================

    async def start_voice_command_session(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> Dict:
        """Start a voice command session."""
        session_id = str(uuid.uuid4())

        vc_session = VoiceCommandSession(
            user_id=user_id,
            session_id=session_id,
            status="active",
            commands=[],
        )
        db.add(vc_session)
        await db.commit()
        await db.refresh(vc_session)

        await self._log_usage(db, user_id, "voice_command_start", 0, 0)

        return {
            "id": str(vc_session.id),
            "session_id": vc_session.session_id,
            "status": vc_session.status,
            "language": vc_session.language,
            "commands_processed": 0,
            "created_at": vc_session.created_at.isoformat(),
        }

    async def process_voice_command(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        session_id: str,
        audio_path: str,
        command_text: Optional[str] = None,
    ) -> Dict:
        """Process voice command in active session."""
        result = await db.execute(
            select(VoiceCommandSession).where(
                and_(
                    VoiceCommandSession.session_id == session_id,
                    VoiceCommandSession.user_id == user_id,
                    VoiceCommandSession.status == "active"
                )
            )
        )
        vc_session = result.scalar_one_or_none()

        if not vc_session:
            raise ValueError("Voice command session not found or inactive")

        if command_text:
            parsed_command = command_text
        elif audio_path:
            stt_result = await self.speech_to_text(db, user_id, audio_path, vc_session.language or "en")
            parsed_command = stt_result.get("text", "")
        else:
            raise ValueError("Either audio_path or command_text is required")

        start_time = time.time()
        prompt = f"""Interpret this voice command and return JSON:
{{
    "action": "action_name",
    "params": {{}},
    "response": "natural language response"
}}

Voice command: {parsed_command}"""

        ai_result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        command_data = {
            "command": parsed_command,
            "action": "unknown",
            "response": ai_result.get("content", "Command processed"),
            "timestamp": datetime.utcnow().isoformat(),
        }

        try:
            parsed = json.loads(ai_result.get("content", "{}"))
            command_data["action"] = parsed.get("action", "unknown")
            command_data["response"] = parsed.get("response", command_data["response"])
        except Exception:
            pass

        commands = vc_session.commands or []
        commands.append(command_data)
        vc_session.commands = commands
        await db.commit()

        await self._log_usage(db, user_id, "voice_command_process", 0, execution_time)

        return {
            "session_id": session_id,
            "command": parsed_command,
            "action": command_data["action"],
            "response": command_data["response"],
            "commands_processed": len(commands),
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # FEATURE 7: AI MEMORY VAULT
    # ===================================================================

    async def backup_memory_vault(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> Dict:
        """Backup user's memory vault."""
        memories = await self.get_user_memories(db, user_id)
        backup_data = json.dumps(memories)

        encrypted_data = self._encrypt(backup_data) if self.cipher else backup_data

        backup = MemoryVaultBackup(
            user_id=user_id,
            backup_data=encrypted_data,
            encrypted=bool(self.cipher),
            memory_count=len(memories),
        )
        db.add(backup)
        await db.commit()
        await db.refresh(backup)

        await self._log_usage(db, user_id, "memory_vault_backup", 0, 0)

        return {
            "id": str(backup.id),
            "memory_count": backup.memory_count,
            "encrypted": backup.encrypted,
            "created_at": backup.created_at.isoformat(),
        }

    async def restore_memory_vault(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        backup_id: uuid.UUID
    ) -> Dict:
        """Restore memory vault from backup."""
        result = await db.execute(
            select(MemoryVaultBackup).where(
                and_(
                    MemoryVaultBackup.id == backup_id,
                    MemoryVaultBackup.user_id == user_id
                )
            )
        )
        backup = result.scalar_one_or_none()

        if not backup:
            raise ValueError("Backup not found")

        try:
            backup_data = self._decrypt(backup.backup_data) if backup.encrypted else backup.backup_data
            memories = json.loads(backup_data)
        except Exception:
            raise ValueError("Failed to parse backup data")

        restored_count = 0
        for mem in memories:
            try:
                await self.save_memory(
                    db=db,
                    user_id=user_id,
                    memory_type=mem.get("type", "context"),
                    key=mem.get("key", ""),
                    value=mem.get("value", ""),
                    importance=mem.get("importance", 5),
                )
                restored_count += 1
            except Exception:
                continue

        await self._log_usage(db, user_id, "memory_vault_restore", 0, 0)

        return {
            "backup_id": str(backup.id),
            "restored_count": restored_count,
            "total_memories": len(memories),
        }

    # ===================================================================
    # FEATURE 8: MULTI-TASK MASTER
    # ===================================================================

    async def execute_multiple_tasks(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        tasks: List[Dict]
    ) -> Dict:
        """Execute multiple tasks in parallel."""
        start_time = time.time()
        batch = TaskBatch(
            user_id=user_id,
            tasks=tasks,
            total_tasks=len(tasks),
            status="running",
            started_at=datetime.utcnow(),
        )
        db.add(batch)
        await db.commit()
        await db.refresh(batch)

        async def run_task(task: Dict) -> Dict:
            task_name = task.get("name", "unnamed")
            task_prompt = task.get("prompt", "")
            try:
                result = await self._call_ollama(task_prompt)
                return {
                    "name": task_name,
                    "status": "completed",
                    "result": result.get("content", "")[:500],
                    "tokens": result.get("tokens", 0),
                }
            except Exception as e:
                return {
                    "name": task_name,
                    "status": "failed",
                    "error": str(e),
                }

        task_results = await asyncio.gather(*[run_task(t) for t in tasks], return_exceptions=True)

        results = []
        completed = 0
        failed = 0

        for res in task_results:
            if isinstance(res, Exception):
                results.append({"name": "unknown", "status": "failed", "error": str(res)})
                failed += 1
            else:
                results.append(res)
                if res.get("status") == "completed":
                    completed += 1
                else:
                    failed += 1

        execution_time = int((time.time() - start_time) * 1000)
        batch.results = results
        batch.completed_tasks = completed
        batch.failed_tasks = failed
        batch.status = "completed"
        batch.completed_at = datetime.utcnow()
        await db.commit()

        await self._log_usage(db, user_id, "multi_task_execution", 0, execution_time)

        return {
            "batch_id": str(batch.id),
            "total_tasks": batch.total_tasks,
            "completed": completed,
            "failed": failed,
            "results": results,
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # FEATURE 9: AI TEACHER MODE
    # ===================================================================

    async def create_course(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        topic: str,
        difficulty: str
    ) -> Dict:
        """Create an AI-generated course."""
        start_time = time.time()

        prompt = f"""Create a course outline for: {topic}
Difficulty: {difficulty}
Output JSON:
{{
    "name": "course_name",
    "description": "description",
    "total_lessons": 10,
    "lessons": [
        {{"number": 1, "title": "lesson_title", "description": "lesson_desc"}}
    ],
    "learning_objectives": ["obj1", "obj2"]
}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            course_data = json.loads(result.get("content", "{}"))
        except Exception:
            course_data = {
                "name": topic,
                "description": f"AI course on {topic}",
                "total_lessons": 5,
                "lessons": [{"number": i, "title": f"Lesson {i}", "description": ""} for i in range(1, 6)],
                "learning_objectives": [],
            }

        course = AICourse(
            user_id=user_id,
            topic=topic,
            difficulty=difficulty,
            description=course_data.get("description", ""),
            course_content=course_data,
            total_lessons=course_data.get("total_lessons", 5),
            model_used="llama3.1",
        )
        db.add(course)
        await db.commit()
        await db.refresh(course)

        await self._log_usage(db, user_id, "course_create", 0, execution_time)

        return {
            "id": str(course.id),
            "topic": course.topic,
            "difficulty": course.difficulty,
            "description": course.description,
            "total_lessons": course.total_lessons,
            "lessons": course_data.get("lessons", []),
            "learning_objectives": course_data.get("learning_objectives", []),
            "created_at": course.created_at.isoformat(),
        }

    async def generate_lesson(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        course_id: uuid.UUID,
        lesson_number: int
    ) -> Dict:
        """Generate a specific lesson for a course."""
        result = await db.execute(
            select(AICourse).where(
                and_(
                    AICourse.id == course_id,
                    AICourse.user_id == user_id
                )
            )
        )
        course = result.scalar_one_or_none()

        if not course:
            raise ValueError("Course not found")

        start_time = time.time()
        prompt = f"""Generate lesson {lesson_number} for the course: {course.topic}
Difficulty: {course.difficulty}
Course content: {json.dumps(course.course_content or {})}

Output JSON:
{{
    "title": "lesson title",
    "content": "detailed lesson content",
    "examples": ["example1", "example2"],
    "exercises": ["exercise1", "exercise2"],
    "quiz": [{{"question": str, "options": [str], "answer": str}}]
}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            lesson_data = json.loads(result.get("content", "{}"))
        except Exception:
            lesson_data = {"title": f"Lesson {lesson_number}", "content": result.get("content", "")}

        course.current_lesson = lesson_number
        course.progress_percent = round((lesson_number / course.total_lessons) * 100, 1)
        if course.progress_percent >= 100:
            course.status = "completed"
            course.completed_at = datetime.utcnow()
        await db.commit()

        await self._log_usage(db, user_id, "lesson_generate", 0, execution_time)

        return {
            "course_id": str(course.id),
            "lesson_number": lesson_number,
            "title": lesson_data.get("title", f"Lesson {lesson_number}"),
            "content": lesson_data.get("content", ""),
            "examples": lesson_data.get("examples", []),
            "exercises": lesson_data.get("exercises", []),
            "quiz": lesson_data.get("quiz", []),
            "progress_percent": course.progress_percent,
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # FEATURE 10: AI BUSINESS ADVISOR
    # ===================================================================

    async def generate_business_plan(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        industry: str,
        budget: str,
        timeline: str
    ) -> Dict:
        """Generate a business plan."""
        start_time = time.time()

        prompt = f"""Generate a comprehensive business plan for:
Industry: {industry}
Budget: {budget}
Timeline: {timeline}

Output JSON:
{{
    "executive_summary": "summary",
    "market_analysis": {{"size": str, "trends": [str], "competitors": [str]}},
    "financial_projections": {{"revenue_year1": str, "revenue_year3": str, "break_even": str}},
    "risk_factors": [str],
    "milestones": [{{"phase": str, "duration": str, "deliverables": [str]}}],
    "team_requirements": [str]
}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            plan_data = json.loads(result.get("content", "{}"))
        except Exception:
            plan_data = {"executive_summary": result.get("content", "")}

        plan = BusinessPlan(
            user_id=user_id,
            industry=industry,
            budget=budget,
            timeline=timeline,
            plan_content=plan_data,
            model_used="llama3.1",
        )
        db.add(plan)
        await db.commit()
        await db.refresh(plan)

        await self._log_usage(db, user_id, "business_plan_generate", 0, execution_time)

        return {
            "id": str(plan.id),
            "industry": plan.industry,
            "budget": plan.budget,
            "timeline": plan.timeline,
            "plan": plan.plan_content,
            "created_at": plan.created_at.isoformat(),
        }

    async def generate_marketing_strategy(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        plan_id: uuid.UUID
    ) -> Dict:
        """Generate marketing strategy for a business plan."""
        result = await db.execute(
            select(BusinessPlan).where(
                and_(
                    BusinessPlan.id == plan_id,
                    BusinessPlan.user_id == user_id
                )
            )
        )
        plan = result.scalar_one_or_none()

        if not plan:
            raise ValueError("Business plan not found")

        start_time = time.time()
        prompt = f"""Generate a marketing strategy for this business plan:
Industry: {plan.industry}
Budget: {plan.budget}
Plan: {json.dumps(plan.plan_content or {})}

Output JSON:
{{
    "target_audience": "description",
    "channels": [str],
    "campaigns": [{{"name": str, "channel": str, "budget": str, "expected_roi": str}}],
    "content_strategy": [str],
    "kpis": [str]
}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        try:
            strategy = json.loads(result.get("content", "{}"))
        except Exception:
            strategy = {"marketing_strategy": result.get("content", "")}

        plan.marketing_strategy = strategy
        plan.status = "marketing_generated"
        await db.commit()

        await self._log_usage(db, user_id, "marketing_strategy_generate", 0, execution_time)

        return {
            "plan_id": str(plan.id),
            "industry": plan.industry,
            "strategy": strategy,
            "execution_time_ms": execution_time,
        }

    # ===================================================================
    # FEATURE 11: UNIVERSAL FORMAT EXPERT
    # ===================================================================

    async def generate_file(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        file_type: str,
        content: str,
        file_format: str
    ) -> Dict:
        """Generate a file in specified format."""
        start_time = time.time()

        mime_map = {
            "json": "application/json",
            "csv": "text/csv",
            "xml": "application/xml",
            "yaml": "text/yaml",
            "markdown": "text/markdown",
            "pdf": "application/pdf",
            "excel": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }

        source = file_type.strip().lower()
        target = file_format.strip().lower()
        converted_content = content

        if source == "json" and target == "csv":
            parsed = json.loads(content)
            rows = parsed if isinstance(parsed, list) else [parsed]
            output = io.StringIO()
            if rows:
                headers = sorted({k for row in rows if isinstance(row, dict) for k in row.keys()})
                writer = csv.DictWriter(output, fieldnames=headers)
                writer.writeheader()
                for row in rows:
                    writer.writerow(row if isinstance(row, dict) else {})
            converted_content = output.getvalue()
        elif source == "csv" and target == "json":
            reader = csv.DictReader(io.StringIO(content))
            converted_content = json.dumps(list(reader), ensure_ascii=False, indent=2)
        elif target in {"txt", "md", "markdown"}:
            converted_content = content

        file_name = f"{source}_{int(time.time())}.{target}"
        file_path = f"./generated_files/{user_id}/{file_name}"

        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(converted_content)

        file_size = os.path.getsize(file_path)

        generated = GeneratedFile(
            user_id=user_id,
            file_type=file_type,
            file_format=file_format,
            file_name=file_name,
            file_path=file_path,
            file_content=converted_content,
            file_size_bytes=file_size,
        )
        db.add(generated)
        await db.commit()
        await db.refresh(generated)

        execution_time = int((time.time() - start_time) * 1000)
        await self._log_usage(db, user_id, "file_generate", 0, execution_time)

        return {
            "id": str(generated.id),
            "file_name": generated.file_name,
            "file_type": generated.file_type,
            "file_format": generated.file_format,
            "file_path": generated.file_path,
            "file_size_bytes": generated.file_size_bytes,
            "mime_type": mime_map.get(file_format, "application/octet-stream"),
            "created_at": generated.created_at.isoformat(),
        }

    async def list_generated_files(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> List[Dict]:
        """List all generated files for a user."""
        result = await db.execute(
            select(GeneratedFile).where(GeneratedFile.user_id == user_id)
            .order_by(desc(GeneratedFile.created_at))
        )
        files = result.scalars().all()

        return [
            {
                "id": str(f.id),
                "file_name": f.file_name,
                "file_type": f.file_type,
                "file_format": f.file_format,
                "file_size_bytes": f.file_size_bytes,
                "download_count": f.download_count,
                "created_at": f.created_at.isoformat(),
            }
            for f in files
        ]

    # ===================================================================
    # FEATURE 12: AI COMPATIBILITY CHECKER
    # ===================================================================

    async def check_compatibility(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        code: str,
        target_platform: str
    ) -> Dict:
        """Check code compatibility with target platform."""
        start_time = time.time()

        prompt = f"""Analyze this code for compatibility with {target_platform}:
{code[:2000]}

Output JSON:
{{
    "compatible": bool,
    "issues": [{{"severity": "error|warning|info", "description": str, "fix": str}}],
    "suggestions": [str],
    "confidence_score": float
}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        compat_data = self._safe_json_loads(result.get("content", "{}"), {
            "compatible": True,
            "issues": [],
            "suggestions": ["Unable to parse analysis"],
            "confidence_score": 0.5,
        })

        # Deterministic heuristic for Python version constraints.
        target_l = target_platform.lower()
        if "python" in target_l:
            code_l = code.lower()
            inferred_min = None
            if "match " in code_l and "case " in code_l:
                inferred_min = "3.10"
            if "tomllib" in code_l:
                inferred_min = "3.11"
            if inferred_min:
                compat_data.setdefault("issues", []).append({
                    "severity": "warning",
                    "description": f"Code appears to require Python {inferred_min}+ features.",
                    "fix": "Use a compatible runtime or refactor syntax/features.",
                })
                compat_data.setdefault("suggestions", []).append(f"Use Python {inferred_min} or newer for best compatibility.")

        check = CompatibilityCheck(
            user_id=user_id,
            code=code[:5000],
            target_platform=target_platform,
            compatible=compat_data.get("compatible", False),
            issues=compat_data.get("issues", []),
            suggestions=compat_data.get("suggestions", []),
            confidence_score=compat_data.get("confidence_score", 0.0),
            model_used="llama3.1",
        )
        db.add(check)
        await db.commit()

        await self._log_usage(db, user_id, "compatibility_check", 0, execution_time)

        return {
            "id": str(check.id),
            "target_platform": check.target_platform,
            "compatible": check.compatible,
            "issues": check.issues,
            "suggestions": check.suggestions,
            "confidence_score": check.confidence_score,
            "execution_time_ms": execution_time,
        }

    async def get_compatibility_history(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> List[Dict]:
        """Get compatibility check history."""
        result = await db.execute(
            select(CompatibilityCheck).where(CompatibilityCheck.user_id == user_id)
            .order_by(desc(CompatibilityCheck.created_at))
        )
        checks = result.scalars().all()

        return [
            {
                "id": str(c.id),
                "target_platform": c.target_platform,
                "compatible": c.compatible,
                "issues_count": len(c.issues),
                "confidence_score": c.confidence_score,
                "created_at": c.created_at.isoformat(),
            }
            for c in checks
        ]

    # ===================================================================
    # FEATURE 13: SMART ROUTER UPGRADE
    # ===================================================================

    async def select_model_for_device(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        task_type: str,
        device_hint: Optional[Dict[str, Any]] = None,
    ) -> Dict:
        """Select optimal model based on device profile."""
        device_result = await db.execute(
            select(DeviceProfile).where(DeviceProfile.user_id == user_id)
        )
        device = device_result.scalar_one_or_none()

        model_map = {
            "text": {"model": "llama3.1", "provider": "ollama", "reason": "Best for text tasks"},
            "code": {"model": "deepseek-r1", "provider": "ollama", "reason": "Best for code tasks"},
            "image": {"model": "stable-diffusion-xl", "provider": "comfyui", "reason": "Best for image tasks"},
            "voice": {"model": "faster-whisper", "provider": "local", "reason": "Best for voice tasks"},
            "document": {"model": "qwen2.5", "provider": "ollama", "reason": "Best for document tasks"},
        }

        selection = model_map.get(task_type, model_map["text"])

        ram_gb = None
        cpu_cores = None
        if device:
            ram_gb = device.ram_gb
            cpu_cores = device.cpu_cores
        if device_hint:
            ram_gb = int(device_hint.get("ram_gb", ram_gb or 0) or 0)
            cpu_cores = int(device_hint.get("cpu_cores", cpu_cores or 0) or 0)

        if ram_gb and ram_gb < 8:
            selection = {
                "model": "llama3:8b",
                "provider": "ollama",
                "reason": "Lightweight model for low-RAM device",
            }
        elif cpu_cores and cpu_cores <= 4 and task_type in {"text", "document"}:
            selection = {
                "model": "mistral",
                "provider": "ollama",
                "reason": "Fast response model for lower CPU devices",
            }

        router_log = ModelRouterLog(
            user_id=user_id,
            task_type=task_type,
            selected_model=selection["model"],
            provider=selection["provider"],
            reason=selection["reason"],
        )
        db.add(router_log)
        await db.commit()

        return selection

    async def get_device_profile(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> Dict:
        """Get user's device profile."""
        result = await db.execute(
            select(DeviceProfile).where(DeviceProfile.user_id == user_id)
        )
        device = result.scalar_one_or_none()

        if not device:
            return {
                "device_name": "default",
                "device_type": "unknown",
                "os": "unknown",
                "cpu_cores": 4,
                "ram_gb": 8,
                "gpu": "none",
                "preferred_model": "llama3.1",
                "capabilities": {},
            }

        return {
            "id": str(device.id),
            "device_name": device.device_name,
            "device_type": device.device_type,
            "os": device.os,
            "cpu_cores": device.cpu_cores,
            "ram_gb": device.ram_gb,
            "gpu": device.gpu,
            "preferred_model": device.preferred_model,
            "capabilities": device.capabilities or {},
            "created_at": device.created_at.isoformat(),
            "updated_at": device.updated_at.isoformat(),
        }

    # ===================================================================
    # FEATURE 14: VOICE CLONING
    # ===================================================================

    async def create_voice_clone(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        voice_name: str,
        audio_sample_path: str,
        consent: bool
    ) -> Dict:
        """Create a voice clone from audio sample."""
        if not consent:
            raise ValueError("Consent is required for voice cloning")
        if not audio_sample_path:
            raise ValueError("A voice sample path is required")

        clone = VoiceClone(
            user_id=user_id,
            voice_name=voice_name,
            audio_sample_path=audio_sample_path,
            consent=consent,
            status="processing",
        )
        db.add(clone)
        await db.commit()
        await db.refresh(clone)

        start_time = time.time()
        prompt = f"""Analyze this voice sample for cloning: {audio_sample_path}
Output JSON: {{"quality": "high|medium|low", "characteristics": [str], "model_recommendation": str}}"""
        ai_result = await self._call_ollama(prompt)

        execution_time = int((time.time() - start_time) * 1000)

        try:
            analysis = json.loads(ai_result.get("content", "{}"))
            clone.accuracy_score = 0.9 if analysis.get("quality") == "high" else 0.7
            clone.voice_metadata = {"characteristics": analysis.get("characteristics", []), "quality": analysis.get("quality", "medium")}
        except Exception:
            clone.accuracy_score = 0.8
            clone.voice_metadata = {"quality": "medium"}

        clone.status = "completed"
        clone.model_used = "piper-tts"
        clone.completed_at = datetime.utcnow()
        await db.commit()

        await self._log_usage(db, user_id, "voice_clone_create", 0, execution_time)

        return {
            "id": str(clone.id),
            "voice_name": clone.voice_name,
            "status": clone.status,
            "accuracy_score": clone.accuracy_score,
            "voice_metadata": clone.voice_metadata,
            "consent": clone.consent,
            "created_at": clone.created_at.isoformat(),
        }

    async def get_voice_clones(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> List[Dict]:
        """Get all voice clones for a user."""
        result = await db.execute(
            select(VoiceClone).where(VoiceClone.user_id == user_id)
            .order_by(desc(VoiceClone.created_at))
        )
        clones = result.scalars().all()

        return [
            {
                "id": str(c.id),
                "voice_name": c.voice_name,
                "status": c.status,
                "accuracy_score": c.accuracy_score,
                "voice_metadata": c.voice_metadata,
                "created_at": c.created_at.isoformat(),
            }
            for c in clones
        ]

    # ===================================================================
    # FEATURE 15: AI NEWS MONITOR
    # ===================================================================

    async def subscribe_to_news(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        topics: List[str],
        frequency: str
    ) -> Dict:
        """Subscribe to news topics."""
        subscription = NewsSubscription(
            user_id=user_id,
            topics=topics,
            frequency=frequency,
            is_active=True,
        )
        db.add(subscription)
        await db.commit()
        await db.refresh(subscription)

        await self._log_usage(db, user_id, "news_subscribe", 0, 0)

        return {
            "id": str(subscription.id),
            "topics": subscription.topics,
            "frequency": subscription.frequency,
            "is_active": subscription.is_active,
            "created_at": subscription.created_at.isoformat(),
        }

    async def generate_news_digest(
        self,
        db: AsyncSession,
        user_id: uuid.UUID
    ) -> Dict:
        """Generate news digest for user's subscriptions."""
        result = await db.execute(
            select(NewsSubscription).where(
                and_(
                    NewsSubscription.user_id == user_id,
                    NewsSubscription.is_active == True
                )
            )
        )
        subscriptions = result.scalars().all()

        if not subscriptions:
            raise ValueError("No active news subscriptions found")

        all_topics = []
        for sub in subscriptions:
            all_topics.extend(sub.topics)

        all_topics = list(set(all_topics))

        start_time = time.time()

        snippets: List[str] = []
        for topic in all_topics[:8]:
            try:
                results = await self._search_searxng(f"latest {topic} security AI news")
                top_items = (results or {}).get("results", [])[:3]
                for item in top_items:
                    snippets.append(
                        f"Topic={topic} | Title={item.get('title', '')} | URL={item.get('url', '')} | Content={item.get('content', '')}"
                    )
            except Exception:
                continue

        prompt = f"""Generate a news digest for these topics: {', '.join(all_topics)}
Include recent developments, key insights, and actionable information.
    Source snippets:\n{chr(10).join(snippets[:24])}

Output JSON:
{{
    "summary": "overall summary",
    "articles": [
        {{"title": str, "topic": str, "summary": str, "importance": "high|medium|low", "source": str}}
    ],
    "key_trends": [str]
}}"""

        result = await self._call_ollama(prompt)
        execution_time = int((time.time() - start_time) * 1000)

        digest_data = self._safe_json_loads(result.get("content", "{}"), {
            "summary": result.get("content", ""),
            "articles": [],
            "key_trends": [],
        })

        active_sub = subscriptions[0] if subscriptions else None

        digest = NewsDigest(
            user_id=user_id,
            subscription_id=active_sub.id if active_sub else None,
            topics=all_topics,
            articles=digest_data.get("articles", []),
            summary=digest_data.get("summary", ""),
        )
        db.add(digest)
        await db.commit()

        await self._log_usage(db, user_id, "news_digest_generate", 0, execution_time)

        return {
            "id": str(digest.id),
            "topics": digest.topics,
            "summary": digest.summary,
            "articles": digest.articles,
            "key_trends": digest_data.get("key_trends", []),
            "generated_at": digest.generated_at.isoformat(),
        }

    # ===================================================================
    # HELPER METHODS
    # ===================================================================

    async def _call_ollama(self, prompt: str, system_prompt: Optional[str] = None) -> Dict:
        """Call self-hosted Ollama."""
        url = f"{self.ollama_url}/api/generate"

        payload = {
            "model": "llama3.1",
            "prompt": prompt,
            "system": system_prompt or "",
            "stream": False,
            "options": {
                "temperature": 0.7,
                "num_predict": 4096,
            },
        }

        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(url, json=payload)
            response.raise_for_status()
            result = response.json()

        return {
            "content": result.get("response", ""),
            "tokens": result.get("eval_count", 0),
        }

    async def _log_usage(
        self,
        db: AsyncSession,
        user_id: uuid.UUID,
        action: str,
        tokens: int,
        execution_time_ms: int
    ):
        """Log feature usage."""
        usage_log = UsageLog(
            user_id=user_id,
            action=action,
            tokens_used=tokens,
            execution_time_ms=execution_time_ms,
        )
        db.add(usage_log)
        await db.commit()


# Singleton instance
advanced_features_service = AdvancedFeaturesService()